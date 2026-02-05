"""
PDF Converter - Convert PDF to various formats with editable text output
"""

from typing import Tuple, Optional, List, Callable
from pathlib import Path
import sys
import os
from src.utilities.logger import get_logger


def get_tesseract_path() -> Optional[str]:
    """Get path to bundled Tesseract executable"""
    # Check if running as PyInstaller bundle
    if getattr(sys, 'frozen', False):
        # Running as compiled executable
        base_path = sys._MEIPASS
    else:
        # Running in development
        base_path = Path(__file__).parent.parent.parent

    # Check for bundled Tesseract
    tesseract_paths = [
        Path(base_path) / 'resources' / 'tesseract' / 'tesseract.exe',
        Path(base_path) / 'tesseract' / 'tesseract.exe',
    ]

    for tess_path in tesseract_paths:
        if tess_path.exists():
            return str(tess_path)

    # Check system PATH
    import shutil
    system_tesseract = shutil.which('tesseract')
    if system_tesseract:
        return system_tesseract

    return None


def get_tesseract_base_dir() -> Optional[str]:
    """Get path to tesseract base directory (parent of tessdata)"""
    # Check if running as PyInstaller bundle
    if getattr(sys, 'frozen', False):
        base_path = sys._MEIPASS
    else:
        base_path = Path(__file__).parent.parent.parent

    # Check for bundled tesseract directory
    tesseract_dirs = [
        Path(base_path) / 'resources' / 'tesseract',
        Path(base_path) / 'tesseract',
    ]

    for tess_dir in tesseract_dirs:
        tessdata_path = tess_dir / 'tessdata'
        if tessdata_path.exists() and (tessdata_path / 'eng.traineddata').exists():
            return str(tess_dir)

    return None


def get_tessdata_dir() -> Optional[str]:
    """Get path to tessdata directory containing traineddata files"""
    base_dir = get_tesseract_base_dir()
    if base_dir:
        tessdata_path = Path(base_dir) / 'tessdata'
        if tessdata_path.exists():
            return str(tessdata_path)
    return None


def setup_tesseract():
    """Configure pytesseract to use bundled Tesseract"""
    try:
        import pytesseract

        tess_path = get_tesseract_path()
        if tess_path:
            pytesseract.pytesseract.tesseract_cmd = tess_path

            # TESSDATA_PREFIX should point to directory containing traineddata files
            # Tesseract looks for: TESSDATA_PREFIX/eng.traineddata
            tess_base_dir = get_tesseract_base_dir()
            if tess_base_dir:
                # Set to tesseract folder where traineddata files are copied
                os.environ['TESSDATA_PREFIX'] = tess_base_dir

            return True
    except ImportError:
        pass
    return False


class PDFConverter:
    """PDF conversion operations - convert PDF to other formats"""

    def __init__(self):
        self.logger = get_logger()
        self.last_error = None
        # Setup bundled Tesseract on initialization
        setup_tesseract()

    def detect_pdf_type(self, pdf_file: str) -> str:
        """
        Detect if PDF contains native text or is scanned (images only)

        Args:
            pdf_file: Path to PDF file

        Returns:
            'text' if PDF has extractable text, 'scanned' if mostly images
        """
        try:
            import fitz

            pdf = fitz.open(pdf_file)
            total_text_length = 0
            total_pages = len(pdf)

            for page_num in range(min(total_pages, 5)):  # Check first 5 pages
                page = pdf[page_num]
                text = page.get_text("text")
                total_text_length += len(text.strip())

            pdf.close()

            # If average text per page is less than 50 chars, likely scanned
            avg_text = total_text_length / min(total_pages, 5)
            return 'text' if avg_text > 50 else 'scanned'

        except Exception as e:
            self.logger.error(f"Error detecting PDF type: {e}")
            return 'text'  # Default to text mode

    def to_word_text_mode(self, pdf_file: str, output_file: str,
                          include_images: bool = True,
                          progress_callback: Optional[Callable] = None) -> Tuple[bool, str]:
        """
        Convert PDF to Word using text extraction (for native/digital PDFs)
        Creates editable Word document with proper text paragraphs

        Args:
            pdf_file: Input PDF file path
            output_file: Output Word file path (.docx)
            include_images: Whether to include images in output
            progress_callback: Callable(current_page, total_pages) for progress

        Returns:
            Tuple of (success: bool, message: str)
        """
        self.last_error = None

        try:
            import fitz
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import io

            # Validate input file
            if not Path(pdf_file).exists():
                self.last_error = f"PDF file not found: {pdf_file}"
                return False, self.last_error

            # Ensure output has .docx extension
            if not output_file.lower().endswith('.docx'):
                output_file += '.docx'

            self.logger.info(f"Converting PDF to Word (Text Mode): {pdf_file}")

            # Open PDF
            pdf = fitz.open(pdf_file)
            total_pages = len(pdf)

            # Create Word document
            doc = Document()

            total_text_extracted = 0

            for page_num in range(total_pages):
                page = pdf[page_num]

                # Progress callback
                if progress_callback:
                    progress_callback(page_num, total_pages)

                # First try simple text extraction (more reliable)
                page_text = page.get_text("text")

                if page_text.strip():
                    # Add text as paragraphs
                    paragraphs = page_text.split('\n')
                    current_para = ""

                    for line in paragraphs:
                        line = line.strip()
                        if line:
                            # Accumulate lines into paragraphs
                            if current_para:
                                current_para += " " + line
                            else:
                                current_para = line
                        else:
                            # Empty line = paragraph break
                            if current_para:
                                para = doc.add_paragraph()
                                run = para.add_run(current_para)
                                run.font.size = Pt(11)
                                total_text_extracted += len(current_para)
                                current_para = ""

                    # Add remaining text
                    if current_para:
                        para = doc.add_paragraph()
                        run = para.add_run(current_para)
                        run.font.size = Pt(11)
                        total_text_extracted += len(current_para)

                # Also try to extract images if requested
                if include_images:
                    try:
                        image_list = page.get_images()
                        for img_index, img in enumerate(image_list):
                            try:
                                xref = img[0]
                                base_image = pdf.extract_image(xref)
                                if base_image:
                                    image_bytes = base_image["image"]
                                    img_stream = io.BytesIO(image_bytes)
                                    # Limit image width to 6 inches
                                    doc.add_picture(img_stream, width=Inches(5.5))
                            except Exception as img_err:
                                self.logger.warning(f"Could not extract image {img_index}: {img_err}")
                    except Exception as img_list_err:
                        self.logger.warning(f"Could not get image list: {img_list_err}")

                # Add page break between pages (except last page)
                if page_num < total_pages - 1:
                    doc.add_page_break()

            pdf.close()

            # Save document
            doc.save(output_file)

            # Final progress
            if progress_callback:
                progress_callback(total_pages, total_pages)

            # Check if any text was extracted
            if total_text_extracted < 100:
                self.logger.warning(f"Very little text extracted ({total_text_extracted} chars). PDF may be scanned.")
                return True, (
                    f"PDF converted to Word.\n\n"
                    f"{total_pages} pages processed.\n"
                    f"⚠️ Very little text was extracted ({total_text_extracted} characters).\n"
                    f"This PDF appears to be scanned. For better results, try OCR mode.\n\n"
                    f"Output: {output_file}"
                )

            self.logger.info(f"PDF converted to Word successfully: {output_file}")
            return True, f"PDF converted to editable Word successfully!\n\n{total_pages} pages converted.\n{total_text_extracted} characters extracted.\nOutput: {output_file}"

        except ImportError as e:
            self.last_error = f"Required library not installed: {str(e)}"
            self.logger.error(self.last_error)
            return False, self.last_error

        except Exception as e:
            self.last_error = f"Error converting PDF to Word: {str(e)}"
            self.logger.error(self.last_error)
            import traceback
            self.logger.error(traceback.format_exc())
            return False, self.last_error

    def to_word_ocr_mode(self, pdf_file: str, output_file: str,
                         language: str = 'eng',
                         progress_callback: Optional[Callable] = None) -> Tuple[bool, str]:
        """
        Convert PDF to Word using OCR (for scanned PDFs)
        Extracts text from images using Tesseract OCR

        Args:
            pdf_file: Input PDF file path
            output_file: Output Word file path (.docx)
            language: OCR language code (e.g., 'eng', 'hin', 'eng+hin')
            progress_callback: Callable(current_page, total_pages) for progress

        Returns:
            Tuple of (success: bool, message: str)
        """
        self.last_error = None

        try:
            import fitz
            from docx import Document
            from docx.shared import Pt
            import pytesseract
            from PIL import Image
            import io

            # Setup bundled Tesseract
            setup_tesseract()

            # Check if Tesseract is available
            try:
                pytesseract.get_tesseract_version()
            except Exception:
                self.last_error = (
                    "Tesseract OCR is not available.\n\n"
                    "The OCR engine could not be initialized.\n"
                    "Please ensure the application was installed correctly."
                )
                return False, self.last_error

            # Validate input file
            if not Path(pdf_file).exists():
                self.last_error = f"PDF file not found: {pdf_file}"
                return False, self.last_error

            # Ensure output has .docx extension
            if not output_file.lower().endswith('.docx'):
                output_file += '.docx'

            self.logger.info(f"Converting PDF to Word (OCR Mode): {pdf_file}")

            # Open PDF
            pdf = fitz.open(pdf_file)
            total_pages = len(pdf)

            # Create Word document
            doc = Document()

            for page_num in range(total_pages):
                page = pdf[page_num]

                # Progress callback
                if progress_callback:
                    progress_callback(page_num, total_pages)

                # Convert page to image
                pix = page.get_pixmap(dpi=300)  # High DPI for better OCR
                img_data = pix.tobytes("png")
                img = Image.open(io.BytesIO(img_data))

                # Run OCR - TESSDATA_PREFIX is set in setup_tesseract()
                try:
                    # Ensure tesseract is configured
                    setup_tesseract()
                    text = pytesseract.image_to_string(img, lang=language)
                except Exception as ocr_error:
                    self.logger.warning(f"OCR error on page {page_num + 1}: {ocr_error}")
                    text = ""

                # Add text to document
                if text.strip():
                    # Split into paragraphs
                    paragraphs = text.split('\n\n')
                    for para_text in paragraphs:
                        if para_text.strip():
                            para = doc.add_paragraph()
                            run = para.add_run(para_text.strip())
                            run.font.size = Pt(11)

                # Add page break between pages (except last page)
                if page_num < total_pages - 1:
                    doc.add_page_break()

            pdf.close()

            # Save document
            doc.save(output_file)

            # Final progress
            if progress_callback:
                progress_callback(total_pages, total_pages)

            self.logger.info(f"PDF converted to Word (OCR) successfully: {output_file}")
            return True, f"PDF converted using OCR successfully!\n\n{total_pages} pages processed.\nOutput: {output_file}"

        except ImportError as e:
            self.last_error = f"Required library not installed: {str(e)}"
            self.logger.error(self.last_error)
            return False, self.last_error

        except Exception as e:
            self.last_error = f"Error converting PDF with OCR: {str(e)}"
            self.logger.error(self.last_error)
            import traceback
            self.logger.error(traceback.format_exc())
            return False, self.last_error

    def to_word_auto(self, pdf_file: str, output_file: str,
                     include_images: bool = True,
                     ocr_language: str = 'eng',
                     progress_callback: Optional[Callable] = None) -> Tuple[bool, str]:
        """
        Auto-detect PDF type and convert accordingly

        Args:
            pdf_file: Input PDF file path
            output_file: Output Word file path (.docx)
            include_images: Whether to include images (for text mode)
            ocr_language: OCR language (for OCR mode)
            progress_callback: Progress callback function

        Returns:
            Tuple of (success: bool, message: str)
        """
        pdf_type = self.detect_pdf_type(pdf_file)
        self.logger.info(f"Detected PDF type: {pdf_type}")

        if pdf_type == 'scanned':
            return self.to_word_ocr_mode(pdf_file, output_file, ocr_language, progress_callback)
        else:
            return self.to_word_text_mode(pdf_file, output_file, include_images, progress_callback)

    def is_tesseract_available(self) -> bool:
        """Check if Tesseract OCR is installed and available"""
        try:
            # First check if bundled Tesseract exists
            tess_path = get_tesseract_path()
            if tess_path and Path(tess_path).exists():
                import pytesseract
                pytesseract.pytesseract.tesseract_cmd = tess_path
                pytesseract.get_tesseract_version()
                return True

            # Try system Tesseract
            import pytesseract
            pytesseract.get_tesseract_version()
            return True
        except Exception:
            return False

    def get_available_ocr_languages(self) -> List[str]:
        """Get list of available OCR languages"""
        try:
            import pytesseract
            langs = pytesseract.get_languages()
            return [l for l in langs if l != 'osd']  # Remove 'osd' (orientation script detection)
        except Exception:
            return ['eng']  # Default to English

    # Keep original methods for backward compatibility
    def to_word(self, pdf_file: str, output_file: str,
                pages: Optional[List[int]] = None) -> Tuple[bool, str]:
        """
        Convert PDF to Word document (.docx) - Legacy method using pdf2docx

        Args:
            pdf_file: Input PDF file path
            output_file: Output Word file path (.docx)
            pages: List of page indices to convert (0-based), None for all pages

        Returns:
            Tuple of (success: bool, message: str)
        """
        # Use new text mode by default
        return self.to_word_text_mode(pdf_file, output_file)

    def to_word_with_progress(self, pdf_file: str, output_file: str,
                              progress_callback=None) -> Tuple[bool, str]:
        """
        Convert PDF to Word with progress callback - Legacy method
        """
        return self.to_word_text_mode(pdf_file, output_file, True, progress_callback)
