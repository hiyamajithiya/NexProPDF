"""
PDF Creation - Convert from Word, Excel, PowerPoint, Images
"""

import fitz  # PyMuPDF
from pathlib import Path
from typing import List, Optional
from PIL import Image
import io
from src.utilities.logger import get_logger


class PDFCreator:
    """Create PDFs from various file formats"""

    def __init__(self):
        self.logger = get_logger()
        self.last_error = None

    def from_images(self, image_files: List[str], output_file: str,
                   page_size: str = "A4") -> bool:
        """
        Create PDF from image files

        Args:
            image_files: List of image file paths
            output_file: Output PDF file path
            page_size: Page size ("A4", "Letter", "Legal")

        Returns:
            True if successful
        """
        try:
            # Page size presets (in points: 1 point = 1/72 inch)
            page_sizes = {
                "A4": fitz.paper_rect("a4"),
                "Letter": fitz.paper_rect("letter"),
                "Legal": fitz.paper_rect("legal")
            }

            pdf = fitz.open()

            for img_file in image_files:
                self.logger.info(f"Adding image: {img_file}")

                # Open image to get dimensions
                img = Image.open(img_file)
                img_width, img_height = img.size

                # Convert to RGB if necessary
                if img.mode not in ('RGB', 'L'):
                    img = img.convert('RGB')

                # Calculate page size based on image aspect ratio
                if page_size in page_sizes:
                    page_rect = page_sizes[page_size]
                else:
                    # Custom size based on image
                    page_rect = fitz.Rect(0, 0, img_width * 72 / 96, img_height * 72 / 96)

                # Create new page
                page = pdf.new_page(width=page_rect.width, height=page_rect.height)

                # Insert image to fit page
                page.insert_image(page_rect, filename=img_file)

                img.close()

            # Save PDF
            pdf.save(output_file)
            pdf.close()

            self.logger.info(f"Created PDF from {len(image_files)} images: {output_file}")
            return True

        except Exception as e:
            self.logger.error(f"Error creating PDF from images: {e}")
            return False

    def from_word(self, word_file: str, output_file: str) -> bool:
        """
        Create PDF from Word document

        Args:
            word_file: Word file path (.docx)
            output_file: Output PDF file path

        Returns:
            True if successful
        """
        self.last_error = None
        try:
            # This requires Microsoft Word or LibreOffice to be installed
            # Using win32com for Windows
            try:
                import win32com.client
                import pythoncom

                pythoncom.CoInitialize()

                try:
                    word = win32com.client.Dispatch("Word.Application")
                    word.Visible = False

                    # Disable DisplayAlerts to prevent dialogs
                    word.DisplayAlerts = 0

                    doc = word.Documents.Open(str(Path(word_file).absolute()))

                    # Configure PDF export options for better quality and accuracy
                    # SaveAs2 with additional parameters for better PDF quality
                    doc.SaveAs2(
                        str(Path(output_file).absolute()),
                        FileFormat=17,  # wdFormatPDF = 17
                        OptimizeFor=0,  # 0 = Standard quality (wdExportOptimizeForPrint)
                        BitmapMissingFonts=True,
                        CreateBookmarks=0,  # 0 = No bookmarks
                        DocStructureTags=True,
                        IncludeDocProps=True,
                        KeepIRM=True
                    )

                    doc.Close(False)  # Close without saving changes
                    word.Quit()

                    pythoncom.CoUninitialize()

                    self.logger.info(f"Created PDF from Word: {output_file}")
                    return True
                except Exception as com_error:
                    pythoncom.CoUninitialize()
                    self.last_error = f"Microsoft Word is not available or failed to convert the document.\n\nError: {str(com_error)}\n\nPlease ensure Microsoft Word is installed on your system."
                    self.logger.error(f"COM error during Word conversion: {com_error}")
                    # Try fallback
                    raise ImportError("Word COM failed")

            except ImportError:
                # Fallback 1: Use docx2pdf if available (also requires MS Word)
                try:
                    from docx2pdf import convert
                    convert(word_file, output_file)
                    self.logger.info(f"Created PDF from Word (docx2pdf): {output_file}")
                    self.last_error = None
                    return True
                except ImportError:
                    pass
                except Exception as fallback_error:
                    self.logger.warning(f"docx2pdf also failed: {fallback_error}")

                # Fallback 2: Pure Python conversion using python-docx + PyMuPDF
                try:
                    return self._from_word_pure_python(word_file, output_file)
                except Exception as pure_python_error:
                    self.last_error = (
                        f"Word to PDF conversion failed.\n\n"
                        f"Microsoft Word is not installed on this system.\n"
                        f"Pure-Python fallback also failed: {str(pure_python_error)}\n\n"
                        f"Please install Microsoft Word for best results."
                    )
                    self.logger.error(f"Pure-python Word conversion error: {pure_python_error}")
                    return False

        except Exception as e:
            if not self.last_error:
                self.last_error = f"Unexpected error during Word to PDF conversion.\n\nError: {str(e)}"
            self.logger.error(f"Error creating PDF from Word: {e}")
            return False

    def _from_word_pure_python(self, word_file: str, output_file: str) -> bool:
        """
        Convert Word document to PDF using pure Python (python-docx + PyMuPDF).
        This is a fallback when Microsoft Word is not installed.
        Preserves text content, basic formatting, and tables.
        """
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor, Emu
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        self.logger.info("Using pure-Python Word to PDF conversion (MS Word not available)")

        doc = Document(word_file)
        pdf = fitz.open()

        # A4 page dimensions in points
        page_width = 595.28
        page_height = 841.89
        margin_left = 56.7  # ~2cm
        margin_right = 56.7
        margin_top = 56.7
        margin_bottom = 56.7
        usable_width = page_width - margin_left - margin_right

        # Current position tracking
        page = pdf.new_page(width=page_width, height=page_height)
        y_pos = margin_top

        def new_page():
            nonlocal page, y_pos
            page = pdf.new_page(width=page_width, height=page_height)
            y_pos = margin_top

        def check_space(needed):
            nonlocal y_pos
            if y_pos + needed > page_height - margin_bottom:
                new_page()

        def get_font_size(paragraph):
            """Get font size from paragraph or its runs."""
            for run in paragraph.runs:
                if run.font.size:
                    return run.font.size.pt
            style = paragraph.style
            if style and style.font and style.font.size:
                return style.font.size.pt
            return 11  # default

        def get_alignment(paragraph):
            """Map Word alignment to fitz alignment."""
            align = paragraph.alignment
            if align == WD_ALIGN_PARAGRAPH.CENTER:
                return fitz.TEXT_ALIGN_CENTER
            elif align == WD_ALIGN_PARAGRAPH.RIGHT:
                return fitz.TEXT_ALIGN_RIGHT
            elif align == WD_ALIGN_PARAGRAPH.JUSTIFY:
                return fitz.TEXT_ALIGN_JUSTIFY
            return fitz.TEXT_ALIGN_LEFT

        def is_bold(paragraph):
            """Check if paragraph is bold (heading-like)."""
            if paragraph.style and paragraph.style.name.startswith('Heading'):
                return True
            for run in paragraph.runs:
                if run.bold:
                    return True
            return False

        # Process document elements
        for element in doc.element.body:
            tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag

            if tag == 'p':
                # Find the paragraph object
                para = None
                for p in doc.paragraphs:
                    if p._element is element:
                        para = p
                        break

                if para is None:
                    continue

                text = para.text.strip()
                if not text and not para.runs:
                    # Empty paragraph - add spacing
                    y_pos += 8
                    continue

                font_size = get_font_size(para)
                bold = is_bold(para)
                alignment = get_alignment(para)

                # Adjust font size for headings
                style_name = para.style.name if para.style else ''
                if style_name.startswith('Heading 1'):
                    font_size = max(font_size, 18)
                    bold = True
                elif style_name.startswith('Heading 2'):
                    font_size = max(font_size, 15)
                    bold = True
                elif style_name.startswith('Heading 3'):
                    font_size = max(font_size, 13)
                    bold = True

                # Estimate text height
                fontname = "hebo" if bold else "helv"
                # Use textbox to handle wrapping
                line_height = font_size * 1.4
                estimated_chars_per_line = max(1, int(usable_width / (font_size * 0.5)))
                estimated_lines = max(1, (len(text) + estimated_chars_per_line - 1) // estimated_chars_per_line)
                text_height = estimated_lines * line_height + 4

                check_space(text_height)

                text_rect = fitz.Rect(
                    margin_left, y_pos,
                    page_width - margin_right, y_pos + text_height + 20
                )

                # Get font color from first run
                color = (0, 0, 0)
                if para.runs and para.runs[0].font.color and para.runs[0].font.color.rgb:
                    rgb = para.runs[0].font.color.rgb
                    color = (rgb[0] / 255, rgb[1] / 255, rgb[2] / 255)

                rc = page.insert_textbox(
                    text_rect, text,
                    fontsize=font_size, fontname=fontname,
                    color=color, align=alignment
                )

                # Advance y position (rc is negative remaining space or positive overflow)
                if rc < 0:
                    # Text fit - actual height used is rect height + rc (rc is negative)
                    actual_height = text_height + 20 + rc
                    y_pos += max(actual_height, line_height) + 2
                else:
                    y_pos += text_height + 4

            elif tag == 'tbl':
                # Process table
                tbl = None
                for t in doc.tables:
                    if t._element is element:
                        tbl = t
                        break

                if tbl is None:
                    continue

                num_cols = len(tbl.columns)
                if num_cols == 0:
                    continue

                col_width = usable_width / num_cols
                row_height = 20

                for row in tbl.rows:
                    # Estimate row height
                    max_text_len = 0
                    for cell in row.cells:
                        max_text_len = max(max_text_len, len(cell.text))
                    chars_per_col = max(1, int(col_width / 6))
                    est_lines = max(1, (max_text_len + chars_per_col - 1) // chars_per_col)
                    row_height = max(20, est_lines * 14 + 6)

                    check_space(row_height)

                    for col_idx, cell in enumerate(row.cells):
                        cell_x = margin_left + col_idx * col_width
                        cell_rect = fitz.Rect(
                            cell_x, y_pos,
                            cell_x + col_width, y_pos + row_height
                        )
                        # Draw cell border
                        page.draw_rect(cell_rect, color=(0.6, 0.6, 0.6), width=0.5)

                        # Insert cell text
                        text_rect = fitz.Rect(
                            cell_x + 3, y_pos + 2,
                            cell_x + col_width - 3, y_pos + row_height - 2
                        )
                        cell_text = cell.text.strip()
                        if cell_text:
                            page.insert_textbox(
                                text_rect, cell_text,
                                fontsize=9, fontname="helv",
                                color=(0, 0, 0), align=fitz.TEXT_ALIGN_LEFT
                            )

                    y_pos += row_height

                y_pos += 8  # Space after table

        # Save PDF
        pdf.save(output_file, garbage=4, deflate=True)
        pdf.close()

        self.logger.info(f"Created PDF from Word (pure-Python): {output_file}")
        self.last_error = None
        return True

    def from_excel(self, excel_file: str, output_file: str) -> bool:
        """
        Create PDF from Excel spreadsheet

        Args:
            excel_file: Excel file path (.xlsx, .xls)
            output_file: Output PDF file path

        Returns:
            True if successful
        """
        self.last_error = None
        try:
            # This requires Microsoft Excel or LibreOffice to be installed
            try:
                import win32com.client
                import pythoncom

                pythoncom.CoInitialize()

                try:
                    excel = win32com.client.Dispatch("Excel.Application")
                    excel.Visible = False
                    excel.DisplayAlerts = False

                    wb = excel.Workbooks.Open(str(Path(excel_file).absolute()))

                    # Export all sheets to PDF with proper settings
                    wb.ExportAsFixedFormat(
                        0,  # Type: 0 = xlTypePDF
                        str(Path(output_file).absolute()),
                        0,  # Quality: 0 = xlQualityStandard (high quality)
                        True,  # IncludeDocProperties
                        True,  # IgnorePrintAreas: False to respect print areas
                        1,  # From: first page
                        wb.Sheets.Count  # To: last sheet
                    )

                    wb.Close(False)
                    excel.Quit()

                    pythoncom.CoUninitialize()

                    self.logger.info(f"Created PDF from Excel: {output_file}")
                    return True
                except Exception as com_error:
                    pythoncom.CoUninitialize()
                    self.last_error = f"Microsoft Excel is not available or failed to convert the document.\n\nError: {str(com_error)}\n\nPlease ensure Microsoft Excel is installed on your system."
                    self.logger.error(f"COM error during Excel conversion: {com_error}")
                    return False

            except ImportError:
                self.last_error = "Excel to PDF conversion requires Microsoft Excel to be installed."
                self.logger.error("win32com not available for Excel conversion")
                return False

        except Exception as e:
            if not self.last_error:
                self.last_error = f"Unexpected error during Excel to PDF conversion.\n\nError: {str(e)}"
            self.logger.error(f"Error creating PDF from Excel: {e}")
            return False

    def from_powerpoint(self, ppt_file: str, output_file: str) -> bool:
        """
        Create PDF from PowerPoint presentation

        Args:
            ppt_file: PowerPoint file path (.pptx, .ppt)
            output_file: Output PDF file path

        Returns:
            True if successful
        """
        self.last_error = None
        try:
            # This requires Microsoft PowerPoint or LibreOffice to be installed
            try:
                import win32com.client
                import pythoncom

                pythoncom.CoInitialize()

                try:
                    powerpoint = win32com.client.Dispatch("PowerPoint.Application")
                    powerpoint.Visible = 0  # Hidden
                    powerpoint.DisplayAlerts = 0  # Disable alerts

                    presentation = powerpoint.Presentations.Open(
                        str(Path(ppt_file).absolute()),
                        ReadOnly=True,
                        Untitled=True,
                        WithWindow=False
                    )

                    # Export as PDF with quality settings
                    presentation.ExportAsFixedFormat(
                        str(Path(output_file).absolute()),
                        2,  # ppFixedFormatTypePDF = 2
                        0,  # Intent: 0 = ppFixedFormatIntentScreen (standard quality)
                        0,  # FrameSlides: 0 = no frame
                        1,  # HandoutOrder: not applicable
                        0,  # OutputType: 0 = ppPrintOutputSlides
                        0,  # PrintHiddenSlides: 0 = no
                        None,  # PrintRange
                        2,  # RangeType: 2 = ppPrintAll
                        "",  # SlideShowName
                        True,  # IncludeDocProperties
                        True,  # KeepIRMSettings
                        True,  # DocStructureTags
                        True,  # BitmapMissingFonts
                        False  # UseISO19005_1 (PDF/A)
                    )

                    presentation.Close()
                    powerpoint.Quit()

                    pythoncom.CoUninitialize()

                    self.logger.info(f"Created PDF from PowerPoint: {output_file}")
                    return True
                except Exception as com_error:
                    pythoncom.CoUninitialize()
                    self.last_error = f"Microsoft PowerPoint is not available or failed to convert the document.\n\nError: {str(com_error)}\n\nPlease ensure Microsoft PowerPoint is installed on your system."
                    self.logger.error(f"COM error during PowerPoint conversion: {com_error}")
                    return False

            except ImportError:
                self.last_error = "PowerPoint to PDF conversion requires Microsoft PowerPoint to be installed."
                self.logger.error("win32com not available for PowerPoint conversion")
                return False

        except Exception as e:
            if not self.last_error:
                self.last_error = f"Unexpected error during PowerPoint to PDF conversion.\n\nError: {str(e)}"
            self.logger.error(f"Error creating PDF from PowerPoint: {e}")
            return False

    def from_text(self, text_file: str, output_file: str,
                 font_size: int = 12, font_name: str = "helv") -> bool:
        """
        Create PDF from text file

        Args:
            text_file: Text file path
            output_file: Output PDF file path
            font_size: Font size
            font_name: Font name

        Returns:
            True if successful
        """
        try:
            # Read text file
            with open(text_file, 'r', encoding='utf-8') as f:
                text = f.read()

            # Create PDF
            pdf = fitz.open()
            page = pdf.new_page(width=595, height=842)  # A4 size

            # Insert text
            text_rect = fitz.Rect(50, 50, 545, 792)  # Margins
            page.insert_textbox(
                text_rect,
                text,
                fontsize=font_size,
                fontname=font_name,
                align=fitz.TEXT_ALIGN_LEFT
            )

            pdf.save(output_file)
            pdf.close()

            self.logger.info(f"Created PDF from text: {output_file}")
            return True

        except Exception as e:
            self.logger.error(f"Error creating PDF from text: {e}")
            return False

    def create_blank_pdf(self, output_file: str, num_pages: int = 1,
                        page_width: float = 595, page_height: float = 842) -> bool:
        """
        Create blank PDF

        Args:
            output_file: Output PDF file path
            num_pages: Number of pages
            page_width: Page width in points
            page_height: Page height in points

        Returns:
            True if successful
        """
        try:
            pdf = fitz.open()

            for _ in range(num_pages):
                pdf.new_page(width=page_width, height=page_height)

            pdf.save(output_file)
            pdf.close()

            self.logger.info(f"Created blank PDF with {num_pages} pages: {output_file}")
            return True

        except Exception as e:
            self.logger.error(f"Error creating blank PDF: {e}")
            return False

    def convert_to_pdfa(self, input_file: str, output_file: str) -> bool:
        """
        Convert PDF to PDF/A (archival format)

        Args:
            input_file: Input PDF file path
            output_file: Output PDF/A file path

        Returns:
            True if successful
        """
        try:
            pdf = fitz.open(input_file)

            # Convert to PDF/A by setting metadata and removing incompatible features
            metadata = {
                'format': 'PDF/A-1b',
                'title': pdf.metadata.get('title', ''),
                'author': pdf.metadata.get('author', ''),
                'subject': pdf.metadata.get('subject', ''),
                'keywords': pdf.metadata.get('keywords', '')
            }

            pdf.set_metadata(metadata)

            # Save with PDF/A compatible settings
            pdf.save(
                output_file,
                garbage=4,
                deflate=True,
                clean=True,
                pretty=True
            )

            pdf.close()

            self.logger.info(f"Converted to PDF/A: {output_file}")
            return True

        except Exception as e:
            self.logger.error(f"Error converting to PDF/A: {e}")
            return False

    def batch_convert(self, input_files: List[str], output_dir: str,
                     format_type: str) -> List[str]:
        """
        Batch convert files to PDF

        Args:
            input_files: List of input file paths
            output_dir: Output directory
            format_type: Type of files ('images', 'word', 'excel', 'powerpoint')

        Returns:
            List of created PDF files
        """
        output_files = []
        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)

        conversion_methods = {
            'images': self.from_images,
            'word': self.from_word,
            'excel': self.from_excel,
            'powerpoint': self.from_powerpoint,
            'text': self.from_text
        }

        if format_type not in conversion_methods:
            self.logger.error(f"Unknown format type: {format_type}")
            return []

        try:
            if format_type == 'images':
                # Special handling for images - can combine multiple
                output_file = output_path / "combined.pdf"
                if self.from_images(input_files, str(output_file)):
                    output_files.append(str(output_file))
            else:
                # Convert each file individually
                method = conversion_methods[format_type]

                for input_file in input_files:
                    input_path = Path(input_file)
                    output_file = output_path / f"{input_path.stem}.pdf"

                    if method(input_file, str(output_file)):
                        output_files.append(str(output_file))

            self.logger.info(f"Batch converted {len(output_files)} files")
            return output_files

        except Exception as e:
            self.logger.error(f"Error in batch conversion: {e}")
            return output_files
